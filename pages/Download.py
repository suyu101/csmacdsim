import streamlit as st
import pandas as pd
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches

# Page config
st.set_page_config(page_title="Download", page_icon="📥", layout="wide")
st.sidebar.page_link('Home.py', label='Home')
st.sidebar.page_link('pages/CSMA_CD.py', label='CSMA/CD')
st.sidebar.page_link('pages/Slotted_Aloha.py', label='Slotted_Aloha')
st.sidebar.markdown("---")

st.title("📥 Download Simulation Reports")
st.info("Generate DOCX reports (with data + graphs) from your latest simulations.")

# --------------------------------------------------------
# Check available results
# --------------------------------------------------------
has_csma = "csma_results" in st.session_state
has_aloha = "aloha_results" in st.session_state

if has_csma or has_aloha:
    sim_choice = st.selectbox(
        "Select Report Type:",
        [
            "CSMA/CD Only" if has_csma else None,
            "Slotted ALOHA Only" if has_aloha else None,
            "Combined Report (CSMA/CD + Slotted ALOHA)" if (has_csma and has_aloha) else None,
        ],
        format_func=lambda x: x or "",
    )

    if st.button("📄 Generate DOCX Report"):
        doc = Document()
        doc.add_heading("Network Protocol Simulation Report", level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        def add_sim_section(doc, name, data_key, plot_key, title):
            data = st.session_state[data_key]
            doc.add_page_break()
            doc.add_heading(f"{name} Simulation", level=1)
            doc.add_heading("Parameters & Metrics", level=2)
            for k, v in data.items():
                if k != "event_log":
                    doc.add_paragraph(f"{k.replace('_',' ').title()}: {v}")
            doc.add_heading("Event Log (First 20 Rows)", level=2)
            df = data["event_log"].head(20)
            t = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = t.rows[0].cells
            for i, c in enumerate(df.columns):
                hdr_cells[i].text = c
            for _, row in df.iterrows():
                row_cells = t.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
            if plot_key in st.session_state:
                doc.add_heading(title, level=2)
                doc.add_picture(st.session_state[plot_key], width=Inches(6))

        if sim_choice == "CSMA/CD Only":
            add_sim_section(doc, "CSMA/CD", "csma_results", "csma_plot", "Timeline")
        elif sim_choice == "Slotted ALOHA Only":
            add_sim_section(doc, "Slotted ALOHA", "aloha_results", "aloha_plot", "Throughput Graph")
        elif sim_choice == "Combined Report (CSMA/CD + Slotted ALOHA)":
            add_sim_section(doc, "CSMA/CD", "csma_results", "csma_plot", "Timeline")
            add_sim_section(doc, "Slotted ALOHA", "aloha_results", "aloha_plot", "Throughput Graph")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        name = "combined_network_report.docx" if "Combined" in sim_choice else f"{sim_choice.lower().replace(' ', '_')}.docx"

        st.download_button(
            label=f"⬇️ Download {sim_choice}",
            data=buffer,
            file_name=name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.warning("⚠️ No simulations found! Run CSMA/CD or Slotted ALOHA first.")
